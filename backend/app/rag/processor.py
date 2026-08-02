import pdfplumber
import os
import re
import uuid
import base64
import docx
from typing import List
from app.vectorstore.chroma_client import chroma_store
from langchain_text_splitters import RecursiveCharacterTextSplitter

# ---------------------------------------------------------------------------
# Section detection (formal IMRaD labelling)
# ---------------------------------------------------------------------------
# Ordered so more specific patterns win. Each entry: (canonical label, regex).
_SECTION_PATTERNS = [
    ("Abstract", r"abstract|summary"),
    ("Introduction", r"introduction|background"),
    ("Related Work", r"related works?|literature review|prior work"),
    ("Methods", r"materials?\s+and\s+methods|methodology|methods?|experimental(\s+(setup|procedures?|section))?|study design|data and methods"),
    ("Results", r"results?(\s+and\s+discussion)?|findings"),
    ("Discussion", r"discussion"),
    ("Conclusion", r"conclusions?|concluding remarks|conclusions? and future work"),
    ("Acknowledgments", r"acknowledge?ments?"),
    ("References", r"references|bibliography|works cited|literature cited"),
]


def detect_sections(text: str):
    """Split plain text into labelled IMRaD-style sections using heading heuristics.
    Returns a list of {"label": str, "text": str}. Works well on cleanly formatted
    papers; anything before the first recognised heading is returned as 'Header'."""
    if not text or not text.strip():
        return []
    lines = text.split("\n")
    sections = []
    current_label = "Header"
    current_lines: List[str] = []

    def _match_header(raw: str):
        s = raw.strip()
        if not s or len(s) > 70:
            return None
        # strip leading numbering like "3.", "3.1", "IV." and trailing colon
        cleaned = re.sub(r"^\s*(\d+(\.\d+)*|[IVXLC]+)[\.\)]\s*", "", s, flags=re.IGNORECASE)
        cleaned = cleaned.strip().rstrip(":.").strip()
        if not cleaned or len(cleaned) > 45:
            return None
        low = cleaned.lower()
        for label, pat in _SECTION_PATTERNS:
            if re.fullmatch(pat, low, flags=re.IGNORECASE):
                return label
        return None

    for line in lines:
        lbl = _match_header(line)
        if lbl:
            if current_lines and any(x.strip() for x in current_lines):
                sections.append({"label": current_label, "text": "\n".join(current_lines).strip()})
            current_label = lbl
            current_lines = []
        else:
            current_lines.append(line)
    if current_lines and any(x.strip() for x in current_lines):
        sections.append({"label": current_label, "text": "\n".join(current_lines).strip()})

    # merge consecutive sections that share a label
    merged = []
    for sec in sections:
        if merged and merged[-1]["label"] == sec["label"]:
            merged[-1]["text"] += "\n\n" + sec["text"]
        else:
            merged.append(sec)
    # if we never found a real heading, don't pretend it's structured
    if len(merged) == 1 and merged[0]["label"] == "Header":
        return []
    return merged


# ---------------------------------------------------------------------------
# JATS / structured XML ingestion
# ---------------------------------------------------------------------------
def _localname(tag: str) -> str:
    return tag.split("}")[-1] if "}" in tag else tag


def _all_text(el) -> str:
    return re.sub(r"\s+", " ", " ".join(el.itertext())).strip()


def parse_jats_xml(file_path: str):
    """Parse a JATS/scholarly XML file into (full_text, sections). JATS already carries
    labelled <sec> elements, so this yields clean Methods/Results/etc. with no ML."""
    import xml.etree.ElementTree as ET
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
    except Exception as e:
        return "", [], f"Could not parse XML: {e}"

    sections = []

    def find_all(local):
        return [el for el in root.iter() if _localname(el.tag) == local]

    # Title
    title = ""
    for el in root.iter():
        if _localname(el.tag) == "article-title":
            title = _all_text(el)
            break
    if title:
        sections.append({"label": "Title", "text": title})

    # Abstract
    for el in find_all("abstract"):
        txt = _all_text(el)
        if txt:
            sections.append({"label": "Abstract", "text": txt})
            break

    # Body sections (top-level <sec> inside <body>)
    body = None
    for el in root.iter():
        if _localname(el.tag) == "body":
            body = el
            break
    if body is not None:
        for sec in list(body):
            if _localname(sec.tag) != "sec":
                continue
            sec_title = ""
            for child in sec:
                if _localname(child.tag) == "title":
                    sec_title = _all_text(child)
                    break
            sec_type = sec.get("sec-type") or ""
            label = (sec_title or sec_type or "Section").strip()
            # normalise to canonical label if it matches a known heading
            low = label.lower().rstrip(":.")
            for canon, pat in _SECTION_PATTERNS:
                if re.fullmatch(pat, low, flags=re.IGNORECASE):
                    label = canon
                    break
            body_text = _all_text(sec)
            if body_text:
                sections.append({"label": label, "text": body_text})

    # References
    refs = []
    for ref in root.iter():
        if _localname(ref.tag) in ("ref", "mixed-citation", "element-citation"):
            t = _all_text(ref)
            if t and len(t) > 10:
                refs.append(t)
    if refs:
        sections.append({"label": "References", "text": "\n".join(refs[:200])})

    full_text = "\n\n".join(f"{s['label']}\n{s['text']}" for s in sections)
    if not full_text.strip():
        # not JATS-shaped; fall back to raw text of the document
        full_text = _all_text(root)
    return full_text, sections, ""


# ---------------------------------------------------------------------------
# OCR fallback for scanned PDFs (Gemini vision — reuses existing GEMINI_API_KEY)
# ---------------------------------------------------------------------------
def _render_pdf_to_pngs(file_path: str, max_pages: int = 15):
    """Render PDF pages to PNG bytes with pypdfium2 (pure-python wheel, no system deps)."""
    try:
        import pypdfium2 as pdfium
    except Exception:
        return []
    out = []
    try:
        pdf = pdfium.PdfDocument(file_path)
        n = min(len(pdf), max_pages)
        for i in range(n):
            page = pdf[i]
            pil = page.render(scale=2.0).to_pil()
            import io
            buf = io.BytesIO()
            pil.save(buf, format="PNG")
            out.append(buf.getvalue())
        pdf.close()
    except Exception as e:
        print(f"pdf render error: {e}")
    return out


def _gemini_ocr(png_list):
    """OCR each page image with Gemini vision. Returns concatenated text ("" on failure)."""
    if not png_list:
        return ""
    try:
        from app.agents.workflow import get_model
        from langchain_core.messages import HumanMessage
    except Exception:
        return ""
    model = get_model()
    if model is None:
        return ""
    pages_text = []
    for idx, png in enumerate(png_list):
        try:
            b64 = base64.b64encode(png).decode()
            msg = HumanMessage(content=[
                {"type": "text", "text": "This is a scanned page from an academic document. Transcribe ALL of its text verbatim, preserving paragraph and heading structure. Output only the transcribed text, nothing else."},
                {"type": "image_url", "image_url": "data:image/png;base64," + b64},
            ])
            resp = model.invoke([msg])
            content = getattr(resp, "content", "")
            if isinstance(content, list):
                content = "".join(p.get("text", "") for p in content if isinstance(p, dict))
            if content and content.strip():
                pages_text.append(content.strip())
        except Exception as e:
            print(f"gemini ocr page {idx} error: {e}")
    return "\n\n".join(pages_text)


class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )

    def extract_structured(self, file_path: str, filename: str):
        """Returns (full_text, sections, ocr_used, error).
        - .xml  -> JATS structured ingestion (labelled sections for free)
        - .pdf  -> text extraction, with a Gemini-vision OCR fallback for scanned files,
                   then heuristic section detection
        - .docx/.txt/.md -> text extraction + heuristic section detection"""
        try:
            ext = filename.lower().split('.')[-1]
            full_text = ""
            sections = []
            ocr_used = False

            if ext == 'xml':
                full_text, sections, err = parse_jats_xml(file_path)
                if not full_text.strip():
                    return "", [], False, (err or "Empty XML document.")
                return full_text, sections, False, ""

            if ext == 'pdf':
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(file_path)
                    for page in reader.pages[:120]:
                        try:
                            t = page.extract_text() or ""
                        except Exception:
                            t = ""
                        if t:
                            full_text += t + "\n\n"
                except Exception:
                    full_text = ""
                if len(full_text.strip()) < 50:
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            full_text = ""
                            for page in pdf.pages[:120]:
                                try:
                                    t = page.extract_text()
                                except Exception:
                                    t = None
                                if t:
                                    full_text += t + "\n\n"
                    except Exception:
                        pass
                # Scanned-PDF fallback: no selectable text -> OCR the page images.
                if len(full_text.strip()) < 50:
                    ocr_text = _gemini_ocr(_render_pdf_to_pngs(file_path))
                    if ocr_text.strip():
                        full_text = ocr_text
                        ocr_used = True
            elif ext == 'docx':
                doc = docx.Document(file_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
            elif ext in ['txt', 'md']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    full_text = f.read()
            else:
                return "", [], False, f"Unsupported extension: {ext}"

            if not full_text.strip():
                return "", [], ocr_used, "Could not extract text from document. It might be a scanned image or empty."

            sections = detect_sections(full_text)
            return full_text, sections, ocr_used, ""
        except Exception as e:
            return "", [], False, str(e)

    def extract_text(self, file_path: str, filename: str) -> tuple[str, str]:
        """Backward-compatible text-only extraction."""
        text, _sections, _ocr, error = self.extract_structured(file_path, filename)
        return text, error

    def process_pdf(self, file_path: str, filename: str) -> tuple[bool, str]:
        """Extracts text from a document, chunks it (tagging each chunk with its
        detected section label), and stores it in ChromaDB."""
        try:
            full_text, sections, _ocr, error = self.extract_structured(file_path, filename)
            if not full_text:
                return False, error

            documents = []
            metadatas = []
            ids = []
            doc_id = str(uuid.uuid4())

            if sections:
                # Chunk within each labelled section so RAG can target Methods/Results/etc.
                i = 0
                for sec in sections:
                    for chunk in self.text_splitter.split_text(sec["text"]):
                        documents.append(chunk)
                        metadatas.append({"source": filename, "chunk": i, "doc_id": doc_id, "section": sec["label"]})
                        ids.append(f"{doc_id}_{i}")
                        i += 1
            else:
                for i, chunk in enumerate(self.text_splitter.split_text(full_text)):
                    documents.append(chunk)
                    metadatas.append({"source": filename, "chunk": i, "doc_id": doc_id, "section": ""})
                    ids.append(f"{doc_id}_{i}")

            if documents:
                chroma_store.add_documents(documents, metadatas, ids)
            return True, ""

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Error processing document {filename}: {e}")
            return False, str(e)

processor = DocumentProcessor()
