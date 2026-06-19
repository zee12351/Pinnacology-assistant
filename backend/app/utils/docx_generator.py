from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import io


def _strip_page_markers(text: str) -> str:
    """Remove AI-inserted page markers like '_Page 1_', '**Page 2**' or a lone 'Page 3' line."""
    text = re.sub(r'(?im)^\s*[_*]{0,2}\s*Page\s*\d+\s*[_*]{0,2}\s*$', '', text)
    text = re.sub(r'(?i)[_*]{1,2}\s*Page\s*\d+\s*[_*]{1,2}', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text


def _add_page_number_footer(doc):
    """Add a real Word page-number field to the footer, aligned bottom-left."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def parse_inline_markdown(paragraph, text):
    """
    Parses **bold** and *italic* markdown tags within a string and adds formatted runs to a paragraph.
    """
    # Simple regex to tokenize string into text, bold, and italic parts
    # It splits the string keeping the delimiters so we can identify them
    # Pattern looks for **text** or *text*
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)

def create_docx_from_markdown(markdown_text: str) -> io.BytesIO:
    """
    Converts markdown text to a highly formatted MS Word Document.
    Returns a BytesIO object containing the .docx file.
    """
    doc = Document()
    
    # Set default font to Times New Roman for an academic/expert look
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Real page numbers in the footer (bottom-left); strip any inline page markers from the text
    _add_page_number_footer(doc)
    markdown_text = _strip_page_markers(markdown_text)

    # Process line by line
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headings
        if line.startswith('# '):
            p = doc.add_heading('', level=1)
            parse_inline_markdown(p, line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading('', level=2)
            parse_inline_markdown(p, line[3:])
        elif line.startswith('### '):
            p = doc.add_heading('', level=3)
            parse_inline_markdown(p, line[4:])
        # Bullet Points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph('', style='List Bullet')
            parse_inline_markdown(p, line[2:])
        # Numbered Lists (basic detection)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph('', style='List Number')
            # Extract the text after the number
            content = re.sub(r'^\d+\.\s', '', line)
            parse_inline_markdown(p, content)
        # Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph('', style='Quote')
            parse_inline_markdown(p, line[2:])
        # Regular Paragraph
        else:
            p = doc.add_paragraph('')
            parse_inline_markdown(p, line)
            
    # Save to a memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
